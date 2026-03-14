"""
Document Export Service

This module provides functionality to export documents in various formats (PDF, DOCX, TXT).
Supports single document and batch export operations with formatting and metadata preservation.
"""

import io
import logging
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from enum import Enum
from datetime import datetime
from pathlib import Path

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    colors = None

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


logger = logging.getLogger(__name__)


class ExportFormat(Enum):
    """Supported export formats"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


@dataclass
class ExportResult:
    """Result of document export operation"""
    success: bool
    format: ExportFormat
    file_size: int
    duration_ms: float
    error_message: Optional[str] = None
    file_path: Optional[str] = None


class DocumentExporter:
    """
    Document exporter supporting multiple formats.

    This service exports documents and their content to PDF, DOCX, or TXT formats
    with proper formatting, metadata, and structure preservation.
    """

    def __init__(self, output_dir: Optional[str] = None):
        """
        Initialize the document exporter.

        Args:
            output_dir: Directory to save exported files. If None, uses current directory.
        """
        self.output_dir = Path(output_dir) if output_dir else Path.cwd()
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Validate library availability
        if not REPORTLAB_AVAILABLE:
            logger.warning("reportlab not installed. PDF export will be disabled.")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not installed. DOCX export will be disabled.")

    def export_document(
        self,
        content: str,
        filename: str,
        export_format: ExportFormat,
        metadata: Optional[Dict[str, Any]] = None,
        title: Optional[str] = None
    ) -> ExportResult:
        """
        Export a single document to the specified format.

        Args:
            content: Document content to export
            filename: Output filename (without extension)
            export_format: Target export format
            metadata: Optional metadata to include in export
            title: Optional document title

        Returns:
            ExportResult with operation status and details

        Raises:
            ValueError: If format is not supported or libraries not available
            IOError: If file writing fails
        """
        start_time = datetime.now()

        # Input validation
        if not isinstance(content, str):
            return ExportResult(
                success=False,
                format=export_format,
                file_size=0,
                duration_ms=0,
                error_message=f"Content must be string, got {type(content).__name__}"
            )

        if not content.strip():
            logger.warning(f"Empty content provided for '{filename}'")

        if not filename or not isinstance(filename, str):
            return ExportResult(
                success=False,
                format=export_format,
                file_size=0,
                duration_ms=0,
                error_message="Filename must be non-empty string"
            )

        # Sanitize filename
        safe_filename = self._sanitize_filename(filename)

        if metadata is not None and not isinstance(metadata, dict):
            return ExportResult(
                success=False,
                format=export_format,
                file_size=0,
                duration_ms=0,
                error_message="Metadata must be dictionary"
            )

        try:
            if export_format == ExportFormat.PDF:
                if not REPORTLAB_AVAILABLE:
                    raise ValueError("PDF export requires reportlab library")
                result = self._export_to_pdf(content, safe_filename, metadata, title)

            elif export_format == ExportFormat.DOCX:
                if not DOCX_AVAILABLE:
                    raise ValueError("DOCX export requires python-docx library")
                result = self._export_to_docx(content, safe_filename, metadata, title)

            elif export_format == ExportFormat.TXT:
                result = self._export_to_txt(content, safe_filename, metadata)

            else:
                raise ValueError(f"Unsupported export format: {export_format}")

            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            result.duration_ms = duration_ms
            result.success = True

            logger.info(
                f"Exported document '{safe_filename}' to {export_format.value} "
                f"({result.file_size} bytes, {duration_ms:.2f}ms)"
            )

            return result

        except ValueError as e:
            # User input error
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            logger.warning(f"Export validation failed for '{safe_filename}': {e}")
            return ExportResult(
                success=False,
                format=export_format,
                file_size=0,
                duration_ms=duration_ms,
                error_message=str(e)
            )
        except (PermissionError, IOError) as e:
            # File system error
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            logger.error(f"File system error exporting '{safe_filename}': {e}")
            return ExportResult(
                success=False,
                format=export_format,
                file_size=0,
                duration_ms=duration_ms,
                error_message=f"Failed to write export file: {str(e)}"
            )
        except Exception as e:
            # Unexpected error
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            logger.error(f"Unexpected error exporting '{safe_filename}': {e}", exc_info=True)
            return ExportResult(
                success=False,
                format=export_format,
                file_size=0,
                duration_ms=duration_ms,
                error_message=f"Unexpected export error: {type(e).__name__}: {str(e)}"
            )

    def export_batch(
        self,
        documents: List[Dict[str, Any]],
        export_format: ExportFormat,
        batch_filename: Optional[str] = None
    ) -> List[ExportResult]:
        """
        Export multiple documents in batch.

        Args:
            documents: List of documents with 'content', 'filename', and optional 'metadata'
            export_format: Target export format
            batch_filename: Optional filename for batch export (for combined formats)

        Returns:
            List of ExportResult objects for each document
        """
        results = []

        for idx, doc in enumerate(documents):
            content = doc.get('content', '')
            filename = doc.get('filename', f'document_{idx + 1}')
            metadata = doc.get('metadata')
            title = doc.get('title')

            result = self.export_document(
                content=content,
                filename=filename,
                export_format=export_format,
                metadata=metadata,
                title=title
            )

            results.append(result)

        success_count = sum(1 for r in results if r.success)
        logger.info(
            f"Batch export complete: {success_count}/{len(documents)} successful, "
            f"format={export_format.value}"
        )

        return results

    def _export_to_pdf(
        self,
        content: str,
        filename: str,
        metadata: Optional[Dict[str, Any]],
        title: Optional[str]
    ) -> ExportResult:
        """Export document to PDF format"""
        output_path = self.output_dir / f"{filename}.pdf"
        buffer = io.BytesIO()

        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Build PDF content
        story = []
        styles = getSampleStyleSheet()

        # Add title if provided
        if title:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.black if colors else None,
                alignment=TA_CENTER,
                spaceAfter=30
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))

        # Add metadata if provided
        if metadata:
            metadata_style = ParagraphStyle(
                'Metadata',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.grey if colors else None
            )
            for key, value in metadata.items():
                metadata_text = f"{key}: {value}"
                story.append(Paragraph(metadata_text, metadata_style))
            story.append(Spacer(1, 12))

        # Add content
        content_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )

        # Split content into paragraphs
        paragraphs = self._split_content_into_paragraphs(content)
        for para in paragraphs:
            # Escape special characters
            escaped_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(escaped_para, content_style))

        # Build PDF
        doc.build(story)

        # Write to file
        pdf_data = buffer.getvalue()
        with open(output_path, 'wb') as f:
            f.write(pdf_data)

        return ExportResult(
            success=True,
            format=ExportFormat.PDF,
            file_size=len(pdf_data),
            duration_ms=0,
            file_path=str(output_path)
        )

    def _export_to_docx(
        self,
        content: str,
        filename: str,
        metadata: Optional[Dict[str, Any]],
        title: Optional[str]
    ) -> ExportResult:
        """Export document to DOCX format"""
        output_path = self.output_dir / f"{filename}.docx"

        # Create document
        doc = Document()

        # Add title if provided
        if title:
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata if provided
        if metadata:
            for key, value in metadata.items():
                doc.add_paragraph(f"{key}: {value}", style='Intense Quote')

            doc.add_paragraph()  # Empty paragraph for spacing

        # Add content
        paragraphs = self._split_content_into_paragraphs(content)
        for para in paragraphs:
            p = doc.add_paragraph(para)
            p.style = 'Normal'

        # Save document
        doc.save(output_path)
        file_size = output_path.stat().st_size

        return ExportResult(
            success=True,
            format=ExportFormat.DOCX,
            file_size=file_size,
            duration_ms=0,
            file_path=str(output_path)
        )

    def _export_to_txt(
        self,
        content: str,
        filename: str,
        metadata: Optional[Dict[str, Any]]
    ) -> ExportResult:
        """Export document to plain text format"""
        output_path = self.output_dir / f"{filename}.txt"

        lines = []

        # Add metadata header if provided
        if metadata:
            lines.append("=" * 50)
            lines.append("DOCUMENT METADATA")
            lines.append("=" * 50)
            for key, value in metadata.items():
                lines.append(f"{key}: {value}")
            lines.append("=" * 50)
            lines.append("")

        # Add content
        lines.append(content)

        # Write to file
        text_content = '\n'.join(lines)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)

        file_size = output_path.stat().st_size

        return ExportResult(
            success=True,
            format=ExportFormat.TXT,
            file_size=file_size,
            duration_ms=0,
            file_path=str(output_path)
        )

    def get_supported_formats(self) -> List[str]:
        """
        Get list of supported export formats.

        Returns:
            List of format names that are available (libraries installed)
        """
        formats = [ExportFormat.TXT.value]  # TXT is always supported

        if REPORTLAB_AVAILABLE:
            formats.append(ExportFormat.PDF.value)

        if DOCX_AVAILABLE:
            formats.append(ExportFormat.DOCX.value)

        return formats

    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename to prevent path traversal and invalid characters.

        Args:
            filename: Original filename

        Returns:
            Sanitized filename safe for file system
        """
        # Remove path traversal attempts
        safe_name = filename.replace('..', '').replace('/', '').replace('\\', '')

        # Remove null bytes
        safe_name = safe_name.replace('\x00', '')

        # Remove leading/trailing spaces and dots
        safe_name = safe_name.strip('. ')

        # Ensure filename is not empty after sanitization
        if not safe_name:
            safe_name = 'document'

        return safe_name

    def _split_content_into_paragraphs(self, content: str) -> List[str]:
        """
        Split content into paragraphs, filtering empty ones.

        Args:
            content: Document content

        Returns:
            List of non-empty paragraph strings
        """
        paragraphs = content.split('\n\n')
        return [p.strip() for p in paragraphs if p.strip()]


# Convenience function for quick exports
def export_document(
    content: str,
    filename: str,
    export_format: str,
    output_dir: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None,
    title: Optional[str] = None
) -> ExportResult:
    """
    Convenience function to export a document.

    Args:
        content: Document content to export
        filename: Output filename (without extension)
        export_format: Target format ('pdf', 'docx', 'txt')
        output_dir: Optional output directory
        metadata: Optional metadata
        title: Optional document title

    Returns:
        ExportResult with operation status
    """
    try:
        format_enum = ExportFormat(export_format.lower())
    except ValueError:
        return ExportResult(
            success=False,
            format=ExportFormat.TXT,
            file_size=0,
            duration_ms=0,
            error_message=f"Invalid format: {export_format}"
        )

    exporter = DocumentExporter(output_dir=output_dir)
    return exporter.export_document(content, filename, format_enum, metadata, title)
